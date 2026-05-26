import os
import pandas as pd
from docx import Document


class AreaDistributionReport:

    AREAS = [
        "Clínica Médica",
        "Cirurgia",
        "Pediatria",
        "Ginecologia e Obstetrícia",
        "Saúde Coletiva"
    ]

    def __init__(self):
        pass

    def padronizar_areas(self, valor):

        texto = str(valor).lower()

        if "clínica" in texto or "clinica" in texto:
            return "Clínica Médica"

        if "cirurgia" in texto:
            return "Cirurgia"

        if "pediatria" in texto:
            return "Pediatria"

        if "gineco" in texto or "obst" in texto:
            return "Ginecologia e Obstetrícia"

        if "coletiva" in texto:
            return "Saúde Coletiva"

        return "Outras"

    def carregar_dados(self):

        caminhos = [
            "outputs/match/compatibilidade_curricular.xlsx",
            "outputs/match/compatibilidade_curricular_top3.xlsx",
            "outputs/semantic/clusters_semanticos.xlsx"
        ]

        caminho_encontrado = None

        for caminho in caminhos:
            if os.path.exists(caminho):
                caminho_encontrado = caminho
                break

        if caminho_encontrado is None:
            raise FileNotFoundError(
                "Nenhum arquivo válido encontrado."
            )

        df = pd.read_excel(caminho_encontrado)

        return df

    def gerar_relatorios(
        self,
        output_resumo="outputs/word/relatorio_5_grandes_areas_por_prova.docx",
        output_anexo="outputs/word/anexo_questoes_por_grande_area.docx"
    ):

        os.makedirs("outputs/word", exist_ok=True)

        df = self.carregar_dados()

        if "grande_area" not in df.columns:
            raise ValueError(
                "Coluna 'grande_area' não encontrada."
            )

        df["grande_area"] = df["grande_area"].apply(
            self.padronizar_areas
        )

        if "tipo_prova" not in df.columns:
            df["tipo_prova"] = "Não informado"

        if "nome_prova" not in df.columns:
            df["nome_prova"] = "Não informado"

        if "numero" not in df.columns:
            df["numero"] = "Não informado"

        if "enunciado" not in df.columns:
            df["enunciado"] = "Não informado"

        self.gerar_resumo(df, output_resumo)

        self.gerar_anexo(df, output_anexo)

        return output_resumo, output_anexo

    def gerar_resumo(self, df, output_path):

        doc = Document()

        doc.add_heading(
            "ATHENA Blueprint — Distribuição das 5 Grandes Áreas",
            level=1
        )

        tabela_base = (
            df.groupby(["tipo_prova", "nome_prova", "grande_area"])
            .size()
            .reset_index(name="n")
        )

        total = (
            df.groupby(["tipo_prova", "nome_prova"])
            .size()
            .reset_index(name="total")
        )

        tabela_base = tabela_base.merge(
            total,
            on=["tipo_prova", "nome_prova"]
        )

        tabela_base["percentual"] = (
            tabela_base["n"] /
            tabela_base["total"] * 100
        ).round(2)

        pivot = tabela_base.pivot_table(
            index=["tipo_prova", "nome_prova"],
            columns="grande_area",
            values="percentual",
            fill_value=0
        ).reset_index()

        tabela = doc.add_table(
            rows=1,
            cols=7
        )

        tabela.style = "Table Grid"

        headers = [
            "Tipo",
            "Prova",
            "Clínica Médica",
            "Cirurgia",
            "Pediatria",
            "Ginecologia e Obstetrícia",
            "Saúde Coletiva"
        ]

        for i, h in enumerate(headers):
            tabela.rows[0].cells[i].text = h

        for _, row in pivot.iterrows():

            linha = tabela.add_row().cells

            linha[0].text = str(row["tipo_prova"])
            linha[1].text = str(row["nome_prova"])

            for i, area in enumerate(self.AREAS, start=2):

                valor = row[area] if area in row else 0

                linha[i].text = f"{round(valor,2)}%"

        doc.save(output_path)

    def gerar_anexo(self, df, output_path):

        doc = Document()

        doc.add_heading(
            "ATHENA Blueprint — Questões por Grande Área",
            level=1
        )

        for area in self.AREAS:

            doc.add_heading(area, level=2)

            subset = df[
                df["grande_area"] == area
            ]

            if subset.empty:
                doc.add_paragraph(
                    "Nenhuma questão encontrada."
                )
                continue

            for _, row in subset.iterrows():

                doc.add_paragraph(
                    f"Questão {row['numero']}"
                )

                doc.add_paragraph(
                    f"Prova: {row['nome_prova']}"
                )

                doc.add_paragraph(
                    f"Tipo: {row['tipo_prova']}"
                )

                doc.add_paragraph(
                    f"Enunciado: {row['enunciado']}"
                )

                doc.add_paragraph("-" * 80)

        doc.save(output_path)
