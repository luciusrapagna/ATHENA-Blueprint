import os
import pandas as pd

from docx import Document
from docx.shared import Inches


class CurricularReportGenerator:

    def __init__(self):
        pass

    def gerar_relatorio(
        self,
        match_path="outputs/match/compatibilidade_curricular_top3.xlsx",
        heatmap_path="outputs/figures/heatmap_curricular.png",
        output_path="outputs/word/relatorio_match_curricular.docx"
    ):

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        df = pd.read_excel(match_path)

        total_questoes = df["questao"].nunique()
        total_aulas = df["aula_compativel"].nunique()

        media_compatibilidade = round(
            df["compatibilidade_percentual"].mean(),
            2
        )

        baixa_aderencia = df[
            df["compatibilidade_percentual"] < 35
        ]

        document = Document()

        document.add_heading(
            "ATHENA Blueprint — Relatório Pedagógico Curricular",
            level=1
        )

        document.add_heading(
            "1. Resumo Executivo",
            level=2
        )

        document.add_paragraph(
            f"Foram analisadas {total_questoes} questões "
            f"e {total_aulas} aulas utilizando embeddings semânticos "
            f"e NLP aplicado ao currículo médico."
        )

        document.add_paragraph(
            f"A média geral de compatibilidade curricular foi de "
            f"{media_compatibilidade}%."
        )

        document.add_heading(
            "2. Questões com Baixa Compatibilidade",
            level=2
        )

        document.add_paragraph(
            f"Foram identificadas {len(baixa_aderencia)} associações "
            f"com compatibilidade inferior a 35%."
        )

        document.add_heading(
            "3. Heatmap Curricular",
            level=2
        )

        if os.path.exists(heatmap_path):
            document.add_picture(
                heatmap_path,
                width=Inches(6.5)
            )

        document.add_heading(
            "4. Recomendação Pedagógica",
            level=2
        )

        document.add_paragraph(
            "O ATHENA Blueprint identificou padrões de aderência "
            "curricular entre questões e aulas, permitindo apoio "
            "ao NDE, revisão curricular longitudinal, análise "
            "de lacunas pedagógicas e alinhamento MEC/DCN."
        )

        document.save(output_path)

        return output_path
