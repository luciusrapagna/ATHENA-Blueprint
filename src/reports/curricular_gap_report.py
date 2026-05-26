import os
import pandas as pd

from docx import Document


class CurricularGapReport:

    def __init__(self):
        pass

    def gerar_relatorio(
        self,
        gaps_path="outputs/analytics/lacunas_curriculares.xlsx",
        output_path="outputs/word/relatorio_lacunas_curriculares.docx"
    ):

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        df = pd.read_excel(gaps_path)

        document = Document()

        document.add_heading(
            "ATHENA Blueprint — Relatório de Lacunas Curriculares",
            level=1
        )

        document.add_heading(
            "1. Resumo Executivo",
            level=2
        )

        document.add_paragraph(
            "O ATHENA Blueprint realizou análise automática "
            "de aderência curricular entre questões e aulas "
            "utilizando embeddings semânticos, NLP médico "
            "e inteligência curricular aplicada."
        )

        document.add_paragraph(
            f"Foram avaliadas {len(df)} aulas no processo "
            f"de detecção automática de lacunas curriculares."
        )

        document.add_heading(
            "2. Diagnóstico Curricular",
            level=2
        )

        tabela = document.add_table(
            rows=1,
            cols=4
        )

        tabela.style = "Table Grid"

        header = tabela.rows[0].cells

        header[0].text = "Aula"
        header[1].text = "Média Compatibilidade"
        header[2].text = "Status Curricular"
        header[3].text = "Recomendação"

        for _, row in df.iterrows():

            linha = tabela.add_row().cells

            linha[0].text = str(row["aula_compativel"])
            linha[1].text = f'{row["media_compatibilidade"]}%'
            linha[2].text = str(row["status_curricular"])
            linha[3].text = str(row["recomendacao_pedagogica"])

        document.add_heading(
            "3. Interpretação Pedagógica",
            level=2
        )

        document.add_paragraph(
            "As aulas classificadas como lacunas curriculares "
            "críticas ou aderência baixa devem ser priorizadas "
            "em revisões do plano pedagógico, alinhamento de "
            "avaliações, ampliação do banco de questões e "
            "fortalecimento da integração longitudinal."
        )

        document.add_heading(
            "4. Aplicabilidade Institucional",
            level=2
        )

        document.add_paragraph(
            "O ATHENA Blueprint fornece suporte estratégico "
            "ao NDE, coordenação do curso e processos MEC/INEP, "
            "permitindo análise longitudinal do currículo médico "
            "e identificação automatizada de fragilidades "
            "pedagógicas."
        )

        document.save(output_path)

        return output_path
